#!/usr/bin/env python3
"""Build a branded Aecon .docx cheat sheet / briefing from scratch.

Usage:
  python3 build_aecon_docx.py

This is a TEMPLATE — copy and modify for specific deliverables.
Key principles:
  - NO nested tables (Word bloats them to 4+ pages)
  - Tight cell margins via tcMar XML
  - Borderless layout tables for side-by-side sections
  - Red badge for section numbers (shaded run background)

Based on: Aecon_GCC_High_Enclave_Cheat_Sheet.docx (commit 1873166)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ═══ BRAND CONSTANTS ═══
RED = RGBColor(0xE5, 0x19, 0x37)
CHARCOAL = RGBColor(0x25, 0x25, 0x25)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MID_GRAY = RGBColor(0x70, 0x70, 0x70)

LOGO = '/data/nextcloud/data/amyn/files/briefings/aecon-assets/logo-aecon-red.png'
OUT = '/data/nextcloud/data/amyn/files/briefings/OUTPUT_FILENAME.docx'

# ═══ HELPER FUNCTIONS ═══

def shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'))

def borders(cell, color="E0E0E0", sz="4", sides="all"):
    opts = {'top': f'<w:top w:val="single" w:sz="{sz}" w:color="{color}"/>',
            'bottom': f'<w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>',
            'left': f'<w:left w:val="single" w:sz="{sz}" w:color="{color}"/>',
            'right': f'<w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'}
    if sides == 'all':
        parts = list(opts.values())
    else:
        parts = [opts[x] for x in sides.split(',')]
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:tcBorders {nsdecls("w")}>{"".join(parts)}</w:tcBorders>'))

def no_borders(cell):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                  '<w:top w:val="none" w:sz="0"/>'
                  '<w:bottom w:val="none" w:sz="0"/>'
                  '<w:left w:val="none" w:sz="0"/>'
                  '<w:right w:val="none" w:sz="0"/>'
                  '</w:tcBorders>'))

def set_margins(cell, top=20, bottom=20, left=40, right=40):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:tcMar {nsdecls("w")}>'
                  f'<w:top w:w="{top}" w:type="dxa"/>'
                  f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
                  f'<w:left w:w="{left}" w:type="dxa"/>'
                  f'<w:right w:w="{right}" w:type="dxa"/>'
                  '</w:tcMar>'))

def rn(p, text, size=9, bold=False, color=CHARCOAL, font='Arial'):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.name = font
    return r

def para(container, space_before=0, space_after=0):
    if hasattr(container, 'add_paragraph'):
        p = container.add_paragraph()
    else:
        p = container.paragraphs[0]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    return p

def badge(p, num):
    r = rn(p, f' {num} ', 8, True, WHITE)
    r._element.get_or_add_rPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E51937"/>'))

def section_heading(doc, num, title):
    p = para(doc, 6, 2)
    badge(p, num)
    rn(p, f'  {title}', 12, True, CHARCOAL)
    p._element.get_or_add_pPr().append(
        parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:color="E0E0E0" w:space="2"/></w:pBdr>'))

def data_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = Inches(widths[i]); c.text = ''
        set_margins(c, 15, 15, 40, 40); shade(c, '252525'); borders(c, '252525', '2')
        rn(c.paragraphs[0], h.upper(), 8, True, WHITE)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.width = Inches(widths[ci]); c.text = ''
            set_margins(c, 10, 10, 40, 40)
            if ri % 2 == 1: shade(c, 'F5F5F5')
            borders(c)
            rn(c.paragraphs[0], val, 8.5)
    return t

def header_block(doc, subtitle, meta):
    ht = doc.add_table(rows=1, cols=2)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc = ht.rows[0].cells[0]; lc.width = Inches(1.2); lc.text = ''; no_borders(lc); set_margins(lc, 0, 0, 0, 0)
    lr = lc.paragraphs[0].add_run(); lr.add_picture(LOGO, width=Inches(1.1))
    tc = ht.rows[0].cells[1]; tc.width = Inches(5.8); tc.text = ''; no_borders(tc); set_margins(tc, 0, 0, 0, 0)
    rn(tc.paragraphs[0], subtitle, 13, True, CHARCOAL)
    mp = tc.add_paragraph(); mp.paragraph_format.space_before = Pt(1)
    rn(mp, meta, 8, color=MID_GRAY)
    for c in ht.rows[0].cells:
        borders(c, 'E51937', '8', 'bottom')

# ═══ BUILD DOCUMENT ═══
# Replace the code below with your specific content.

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.4)
    s.bottom_margin = Inches(0.4)
    s.left_margin = Inches(0.5)
    s.right_margin = Inches(0.5)

ns = doc.styles['Normal']
ns.font.name = 'Arial'; ns.font.size = Pt(9); ns.font.color.rgb = CHARCOAL
ns.paragraph_format.space_before = Pt(0); ns.paragraph_format.space_after = Pt(0)
ns.paragraph_format.line_spacing = 1.05

# ... Add your content here ...

doc.save(OUT)
print(f'Saved: {OUT}')
