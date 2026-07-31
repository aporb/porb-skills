#!/usr/bin/env python3
"""Convert markdown proposal artifacts to properly formatted .docx for federal submissions.

Handles: headings (h1-h3), paragraphs with inline bold/italic, markdown tables → Word tables,
horizontal rules, bullet points, numbered lists. Sets margins, fonts, and spacing per NOFO specs.

Usage: modify the __main__ block to point at your source .md files and desired output paths.

Key NOFO format rules handled:
- Font: Open Sans 15pt body, Calibri 10pt tables (customizable via params)
- Margins: 1 inch all sides (customizable)
- Single-spaced with light paragraph spacing
- Tables: Table Grid style, header row bold, compact cell spacing

Dependencies: python-docx (pip install python-docx)
"""

import re, sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_docx(md_path, docx_path, body_font='Open Sans', body_size=15, header_font='Open Sans', 
                table_font='Calibri', table_size=10, margins=Inches(1)):
    """Convert a markdown file to a properly formatted .docx file.
    
    Args:
        md_path: Path to source markdown file
        docx_path: Output .docx path
        body_font: Font for body text (default: Open Sans for State Dept NOFOs)
        body_size: Point size for body text (default: 15 for State Dept NOFOs)
        header_font: Font for headings
        table_font: Font for table cells (default: Calibri 12 for budget tables)
        table_size: Point size for table cells
        margins: docx.shared.Inches value for all 4 margins
    """
    
    with open(md_path, 'r') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = body_font
    font.size = Pt(body_size)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = margins
        section.bottom_margin = margins
        section.left_margin = margins
        section.right_margin = margins
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Detect tables (lines starting with |)
        if line.strip().startswith('|') and '---' not in line:
            in_table = True
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # Filter out separator rows
            data_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if all(c.replace('-','').replace(':','').replace(' ','') == '' for c in cells):
                    continue  # Skip separator rows
                data_rows.append(cells)
            
            if data_rows:
                num_cols = max(len(row) for row in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for r_idx, row in enumerate(data_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            run = p.add_run(cell_text)
                            run.font.name = table_font
                            run.font.size = Pt(table_size)
                            # Bold for header row
                            if r_idx == 0:
                                run.bold = True
                            # Remove cell margins
                            p.paragraph_format.space_before = Pt(1)
                            p.paragraph_format.space_after = Pt(1)
                
                doc.add_paragraph()  # Space after table
            continue
        
        # Headers
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(level=1)
            run = p.add_run(line[2:].strip())
            run.font.name = header_font
            i += 1
            continue
        
        if line.startswith('## ') and not line.startswith('### '):
            p = doc.add_heading(level=2)
            run = p.add_run(line[3:].strip())
            run.font.name = header_font
            i += 1
            continue
        
        if line.startswith('### '):
            p = doc.add_heading(level=3)
            run = p.add_run(line[4:].strip())
            run.font.name = header_font
            i += 1
            continue
        
        # Horizontal rules
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Handle bold in bullets
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text, body_font, body_size)
            i += 1
            continue
        
        # Numbered lists
        num_match = re.match(r'^(\d+)[.\)]\s+(.*)', line.strip())
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text, body_font, body_size)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        process_inline_formatting(p, line.strip(), body_font, body_size)
        i += 1
    
    doc.save(docx_path)
    return True

def process_inline_formatting(paragraph, text, font_name, font_size):
    """Process inline bold (**text**) and italic (*text*) formatting."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
        else:
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size)

if __name__ == '__main__':
    # EXAMPLE USAGE — modify paths for your submission:
    # Proven in DFOP0018157 pipeline (July 2026): 3 artifacts → 3 .docx files
    base = os.path.dirname(os.path.abspath(__file__))
    
    conversions = [
        ('artifacts/proposal-narrative.md',   'final-package/Proposal_Narrative.docx'),
        ('artifacts/budget-package.md',       'final-package/Budget_Package.docx'),
        ('artifacts/supporting-documents.md', 'final-package/Supporting_Documents.docx'),
    ]
    
    for src_rel, dst_rel in conversions:
        src = os.path.join(base, src_rel)
        dst = os.path.join(base, dst_rel)
        if os.path.exists(src):
            print(f"Converting {src_rel}...")
            create_docx(src, dst, body_font='Open Sans', body_size=15, table_font='Calibri', table_size=10)
            print(f"  ✅ Done ({os.path.getsize(dst)/1024:.1f} KB)")
        else:
            print(f"  ⚠️ Source not found: {src}")
